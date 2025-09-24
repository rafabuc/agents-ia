import os
from typing import List, Dict, Any
from pathlib import Path
from datetime import datetime

# Import with fallback
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from langchain.schema import Document
    LANGCHAIN_AVAILABLE = True
except ImportError:
    # Fallback Document class
    class Document:
        def __init__(self, page_content, metadata=None):
            self.page_content = page_content
            self.metadata = metadata or {}
    LANGCHAIN_AVAILABLE = False

from config.settings import settings
from storage.database_manager import DatabaseManager
from utils.logger import logger


class DocumentProcessor:
    """Processes various document types for RAG ingestion."""
    
    def __init__(self):
        self.db_manager = DatabaseManager()
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.md'}
        
        # Log available processors
        available = []
        if PDF_AVAILABLE:
            available.append("PDF")
        if DOCX_AVAILABLE:
            available.append("DOCX")
        available.extend(["TXT", "MD"])
        
        logger.info(f"DocumentProcessor initialized with support for: {', '.join(available)}")
    
    def process_directory(self, directory_path: str, 
                         source_type: str = "knowledge_base") -> List[Document]:
        """Process all supported documents in a directory."""
        documents = []
        
        if not os.path.exists(directory_path):
            logger.warning(f"Directory does not exist: {directory_path}")
            return documents
        
        for file_path in Path(directory_path).rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc = self.process_file(str(file_path), source_type)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    logger.error(f"Error processing file {file_path}: {str(e)}")
        
        logger.info(f"Processed {len(documents)} documents from {directory_path}")
        return documents
    
    def process_file(self, file_path: str, source_type: str = "knowledge_base") -> Document:
        """Process a single file based on its extension."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File does not exist: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                content = self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                content = self._extract_docx_text(file_path)
            elif file_extension in ['.txt', '.md']:
                content = self._extract_text_file(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Create document object
            document = Document(
                page_content=content,
                metadata={
                    "source": file_path,
                    "filename": os.path.basename(file_path),
                    "file_type": file_extension,
                    "source_type": source_type,
                    "file_size": os.path.getsize(file_path),
                    "processed_at": datetime.utcnow().isoformat()
                }
            )
            
            # Store document info in database
            self._store_document_info(document, file_path, source_type)
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")
        
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        text += f"\n--- Page {page_num + 1} ---\n[Text extraction failed]\n"
                        
        except Exception as e:
            logger.error(f"Error extracting PDF text from {file_path}: {str(e)}")
            raise
        
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        text = ""
        
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
                    
        except Exception as e:
            logger.error(f"Error extracting DOCX text from {file_path}: {str(e)}")
            raise
        
        return text.strip()
    
    def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text or markdown file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    return content.strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                raw_content = file.read()
                content = raw_content.decode('utf-8', errors='ignore')
                return content.strip()
                
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            raise
    
    def _store_document_info(self, document: Document, file_path: str, source_type: str):
        """Store document information in database."""
        try:
            self.db_manager.create_document(
                filename=document.metadata["filename"],
                file_path=file_path,
                document_type=document.metadata["file_type"],
                source=source_type
            )
            logger.debug(f"Stored document info in database: {document.metadata['filename']}")
        except Exception as e:
            logger.warning(f"Could not store document info in database: {str(e)}")
    
    def process_text_content(self, content: str, filename: str = "text_content", 
                           source_type: str = "manual") -> Document:
        """Process raw text content into a document."""
        try:
            document = Document(
                page_content=content,
                metadata={
                    "source": f"text_input_{filename}",
                    "filename": filename,
                    "file_type": ".txt",
                    "source_type": source_type,
                    "file_size": len(content.encode('utf-8')),
                    "processed_at": datetime.utcnow().isoformat()
                }
            )
            
            logger.info(f"Processed text content: {filename}")
            return document
            
        except Exception as e:
            logger.error(f"Error processing text content: {str(e)}")
            raise
    
    def get_supported_formats(self) -> Dict[str, bool]:
        """Get list of supported file formats and their availability."""
        return {
            "pdf": PDF_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "txt": True,
            "md": True
        }
    
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """Validate if a file can be processed."""
        if not os.path.exists(file_path):
            return {
                "valid": False,
                "error": "File does not exist"
            }
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_extensions:
            return {
                "valid": False,
                "error": f"Unsupported file type: {file_extension}"
            }
        
        # Check if required dependencies are available
        if file_extension == '.pdf' and not PDF_AVAILABLE:
            return {
                "valid": False,
                "error": "PyPDF2 not available for PDF processing"
            }
        
        if file_extension == '.docx' and not DOCX_AVAILABLE:
            return {
                "valid": False,
                "error": "python-docx not available for DOCX processing"
            }
        
        # Check file size (optional - set reasonable limit)
        try:
            file_size = os.path.getsize(file_path)
            max_size = 50 * 1024 * 1024  # 50MB limit
            
            if file_size > max_size:
                return {
                    "valid": False,
                    "error": f"File too large: {file_size / (1024*1024):.1f}MB (max: 50MB)"
                }
        except Exception as e:
            return {
                "valid": False,
                "error": f"Cannot access file: {str(e)}"
            }
        
        return {
            "valid": True,
            "file_type": file_extension,
            "file_size": file_size
        }